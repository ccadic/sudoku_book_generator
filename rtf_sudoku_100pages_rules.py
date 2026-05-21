import numpy as np
import random
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_UNDERLINE
import os

def print_grid(grid):
    """Displays the Sudoku grid in a readable format."""
    for row in grid:
        print(" ".join(str(num) if num != 0 else "." for num in row))

def is_valid(grid, row, col, num):
    """Checks if a number can be placed in a given cell."""
    if num in grid[row]:
        return False
    if num in grid[:, col]:
        return False
    start_row, start_col = 3 * (row // 3), 3 * (col // 3)
    if num in grid[start_row:start_row + 3, start_col:start_col + 3]:
        return False
    return True

def solve(grid):
    """Solves the Sudoku puzzle using backtracking."""
    for row in range(9):
        for col in range(9):
            if grid[row, col] == 0:
                for num in range(1, 10):
                    if is_valid(grid, row, col, num):
                        grid[row, col] = num
                        if solve(grid):
                            return True
                        grid[row, col] = 0
                return False
    return True

def generate_complete_grid():
    """Generates a complete Sudoku grid."""
    grid = np.zeros((9, 9), dtype=int)
    for row in range(9):
        for col in range(9):
            nums = list(range(1, 10))
            random.shuffle(nums)
            for num in nums:
                if is_valid(grid, row, col, num):
                    grid[row, col] = num
                    if not solve(grid):
                        grid[row, col] = 0
                        continue
                    break
    return grid

def remove_numbers(grid, difficulty):
    """Removes numbers from a completed grid to create a puzzle."""
    levels = {"easy": 35, "medium": 45, "hard": 55}
    cells_to_remove = levels.get(difficulty, 45)

    grid_copy = grid.copy()
    while cells_to_remove > 0:
        row, col = random.randint(0, 8), random.randint(0, 8)
        if grid_copy[row, col] != 0:
            grid_copy[row, col] = 0
            cells_to_remove -= 1

    return grid_copy

def draw_sudoku(grid, difficulty, cell_size=50):
    """Draws a Sudoku grid and returns it as a PIL Image."""
    grid_size = 9 * cell_size
    img = Image.new("RGB", (grid_size, grid_size + 30), "white")  # Extra space for difficulty label
    draw = ImageDraw.Draw(img)
    font = ImageFont.truetype("arial.ttf", size=20)
    label_font = ImageFont.truetype("arial.ttf", size=16)

    # Draw grid lines
    for i in range(10):
        line_width = 3 if i % 3 == 0 else 1
        draw.line([(i * cell_size, 0), (i * cell_size, grid_size)], fill="black", width=line_width)
        draw.line([(0, i * cell_size), (grid_size, i * cell_size)], fill="black", width=line_width)

    # Add numbers
    for row in range(9):
        for col in range(9):
            num = grid[row, col]
            if num != 0:
                num_text = str(num)
                text_bbox = draw.textbbox((0, 0), num_text, font=font)
                text_width = text_bbox[2] - text_bbox[0]
                text_height = text_bbox[3] - text_bbox[1]
                x = col * cell_size + (cell_size - text_width) // 2
                y = row * cell_size + (cell_size - text_height) // 2
                draw.text((x, y), num_text, fill="black", font=font)

    # Add difficulty label
    label_text = difficulty.capitalize()
    label_bbox = draw.textbbox((0, 0), label_text, font=label_font)
    label_width = label_bbox[2] - label_bbox[0]
    draw.text(((grid_size - label_width) // 2, grid_size + 5), label_text, fill="black", font=label_font)

    return img

def create_full_page_image(grids, difficulties, output_path, page_size=(15.24, 22.86), dpi=210):
    """Creates a full page with 3x4 Sudoku grids and saves as an image."""
    page_width = int(page_size[0] * dpi / 2.54)
    page_height = int(page_size[1] * dpi / 2.54)
    grid_size = min(page_width // 3, (page_height - 120) // 4) - 10

    page = Image.new("RGB", (page_width, page_height), "white")

    for i, (grid, difficulty) in enumerate(zip(grids, difficulties)):
        row, col = divmod(i, 3)
        x_offset = col * grid_size
        y_offset = row * grid_size + row * 30  # Add 30 pixels between rows

        sudoku_img = draw_sudoku(grid, difficulty, cell_size=grid_size // 9)
        page.paste(sudoku_img, (x_offset, y_offset))

    page.save(output_path, dpi=(dpi, dpi))

def create_word_document_with_pages(output_doc, total_pages):
    """Creates a Word document with multiple pages, including title, rules, and Sudoku grids."""
    doc = Document()

    # Add title page
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_paragraph.add_run("LE SUDOKU 2025")
    title_run.bold = True
    title_run.font.size = Pt(36)
    doc.add_page_break()

    # Add rules page
    rules = [
        ("Règle du jeu du Sudoku pour débutants", True, True),
        ("Le Sudoku est un jeu de réflexion et de logique qui consiste à remplir une grille de chiffres en respectant certaines règles.", False, False),
        ("Objectif", True, True),
        ("Le but du Sudoku est de remplir une grille de 9x9 cases avec des chiffres allant de 1 à 9, de manière à ce que :\n", False, False),
        ("- Chaque chiffre apparaisse une seule fois dans chaque ligne.\n", False, False),
        ("- Chaque chiffre apparaisse une seule fois dans chaque colonne.\n", False, False),
        ("- Chaque chiffre apparaisse une seule fois dans chaque sous-grille de 3x3 cases (appelée \"région\" ou \"bloc\").\n", False, False),
        ("La grille", True, True),
        ("La grille complète est composée de :\n", False, False),
        ("- 9 lignes horizontales.\n", False, False),
        ("- 9 colonnes verticales.\n", False, False),
        ("- 9 sous-grilles de 3x3 cases (chaque sous-grille est délimitée par des lignes épaisses).\n", False, False),
        ("Certaines cases de la grille sont pré-remplies au début du jeu. Ces chiffres ne peuvent pas être modifiés et servent de point de départ pour résoudre le puzzle.\n", False, False),
        ("Règles en détail", True, True),
        ("- Un chiffre par ligne :\n  Dans chaque ligne horizontale, les chiffres de 1 à 9 doivent apparaître une seule fois. Aucun chiffre ne peut être répété dans une même ligne.\n", False, False),
        ("- Un chiffre par colonne :\n  Dans chaque colonne verticale, les chiffres de 1 à 9 doivent également apparaître une seule fois. Aucun chiffre ne peut être répété dans une même colonne.\n", False, False),
        ("- Un chiffre par sous-grille :\n  Chaque sous-grille de 3x3 cases doit contenir les chiffres de 1 à 9, sans répétition.\n", False, False),
        ("Comment jouer ?", True, True),
        ("1. Comprendre les indices\n  Commencez par observer la grille et repérez les chiffres qui sont déjà remplis.\n", False, False),
        ("2. Utiliser la logique\n  Analysez ligne par ligne, colonne par colonne, et les sous-grilles.\n", False, False),
        ("3. Remplir progressivement\n", False, False),
        ("4. Vérifiez vos placements\n", False, False),
    ]

    for text, bold, underline in rules:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = bold
        if underline:
            run.underline = WD_UNDERLINE.SINGLE
    doc.add_page_break()

    for page_number in range(1, total_pages + 1):
        difficulties = ["easy"] * 4 + ["medium"] * 4 + ["hard"] * 4
        grids = [remove_numbers(generate_complete_grid(), difficulty) for difficulty in difficulties]

        image_path = f"sudoku_page_{page_number}.png"
        create_full_page_image(grids, difficulties, image_path)

        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(6), height=Inches(8))

        if page_number != total_pages:
            doc.add_page_break()

        os.remove(image_path)

    doc.save(output_doc)

def main():
    total_pages = 100
    output_doc = "Sudoku_Book.docx"
    create_word_document_with_pages(output_doc, total_pages)
    print(f"Sudoku book with {total_pages} pages saved as {output_doc}")

if __name__ == "__main__":
    main()
